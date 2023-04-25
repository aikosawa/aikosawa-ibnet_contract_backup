"""
ドキュメント生成のロジックを持つモジュール
"""


from app import docx_helper, replace
from app import time_helper, jikkin_sheet
from app import xl_helper
from app.config import Config
from app.model import JointGuarantor, Product, ProductInput
from app.time_helper import strftime
from copy import copy
from dataclasses import dataclass
from docx import Document
from logging import getLogger, INFO
from openpyxl.cell.cell import Cell
from openpyxl.styles.borders import Border, Side
from openpyxl.styles.fills import PatternFill
from openpyxl.worksheet.worksheet import Worksheet
from os import times
from typing import Any, ChainMap, Iterable, List, Optional, Tuple, cast
import win32print
import win32api
import win32com.client as win32

from collections import ChainMap

import functools
import itertools
import more_itertools
import openpyxl
import os
import platform
import re
import subprocess
import sys
import copy

logger = getLogger(__name__)

output_root_form_no = {3, 15, 16, 17, 18}

before_ext = re.compile(r'(?=\.(docx|xlsx))')

chacot_flg = [0, 0, 0, 0]


def word_to_pdf_2_pages_per_sheet(input_file, output_file):
    word = win32.gencache.EnsureDispatch('Word.Application')
    word.Visible = False
    word.DisplayAlerts = False

    # デフォルトプリンタ名を取得
    default_printer = win32print.GetDefaultPrinter()

    input_file = os.path.abspath(input_file)
    output_file = os.path.abspath(output_file)

    # Word文書を開く
    doc = word.Documents.Open(input_file)

    if "【集約印刷】" in input_file:
        # 1枚に2ページ単位で印刷する設定
        word.ActivePrinter = default_printer
        word.PrintOut(
            OutputFileName=output_file,
            Item=win32.constants.wdPrintDocumentContent,
            Copies=1,
            Pages="1-2",
            Collate=True,
            Background=False,
            PrintToFile=True,
            Range=win32.constants.wdPrintAllDocument,
            ManualDuplexPrint=False,
            PrintZoomColumn=2,
            PrintZoomRow=1,
            PrintZoomPaperWidth=0,
            PrintZoomPaperHeight=0
            )
    else:
        # 通常のPDF化を行う
        doc.SaveAs(output_file, FileFormat=win32.constants.wdFormatPDF)

    # Word文書を閉じる
    doc.Close(SaveChanges=False)
    word.Quit()


def excel_to_pdf(excel_path, pdf_path):
    """
    ExcelからPDFに変換する関数
    """

    excel = win32.gencache.EnsureDispatch('Excel.Application')
    # excel = win32.Dispatch("Excel.Application")
    excel.Visible = False
    excel.DisplayAlerts = False
    # xlTypePDF定数を取得
    xlTypePDF = win32.constants.xlTypePDF

    try:
        file = excel.Workbooks.Open(os.path.abspath(excel_path))
        file.Worksheets(1).Select()
        file.ActiveSheet.ExportAsFixedFormat(xlTypePDF, os.path.abspath(pdf_path))
    except Exception as e:
        print(f"Failed to convert {excel_path} to PDF: {e}")
    finally:
        file.Close()
        excel.Application.Quit()


def logging_output(ipt: str, opt: str, level=INFO):
    """
    生成された帳票が何をもとにして何を出力したかをロギングする関数
    """
    logger.log(level, f'  出力{ipt}    ->    {opt}')


@dataclass(frozen=True)
class ChohyoGenerator:
    """
    英語のよい名前が思いつかなかったので、帳票生成器と名付けました。

    NOTE: ファイル名は各関数でうまく名前をつけること
    """

    config: Config
    template_root_path: str
    output_root_path: str

    def _wait_all_updates(self, file_iterable: Iterable[str], throttle_sec: float = 1):
        """
        すべてのファイルが更新されることを確認するためのロジック及び表示。
        excelの計算をpythonで実行できないので実装してある。


        FIXME: 開発で使ったマシン上では、開いたタイミングでst_mtimeが更新されてしまった
        つまり上書き保存すると2回st_mtimeが変わるはずなのでそれらを変数に持てばこの関数も
        意図通りに動くはず
        """

        files = list(file_iterable)

        def updated(fpath: str):
            """
            上書き保存がされていたらTrue
            """
            stat = os.stat(fpath)
            return stat.st_ctime != stat.st_mtime

        def refresh():
            return [[fpath, updated(fpath)] for fpath in files]

        state = refresh()

        while any(not updated for _, updated in state):
            logger.info('すべてのファイルをExcelで開き、上書き保存を実行してください')
            logger.info('--------------------------------------------------')
            logger.info('\t状態\tファイル名')
            for fpath, valid in state:
                logger.info(f'\t{"確認済" if  valid else "未確認"}\t{fpath}')
            logger.info('--------------------------------------------------')

            while state == (new_state := refresh()):
                times.sleep(throttle_sec)

            state = new_state

    def gen_all_doc(self, product_inputs: List[ProductInput]):
        """
        帳票出力を行う
        """

        assert len(product_inputs) >= 1

        def jikkin_path(product_input: ProductInput) -> Tuple[str, str]:

            if not isinstance(product_input, ProductInput):
                raise TypeError()
            form_no, filename = more_itertools.first_true(
                self.config.get_product_doc_info(product_input).items(),
                default=(None, None),
                pred=lambda x: (x[0] in {7, 8, 9} and x[1] is not None)
            )
            if filename is None:
                raise RuntimeError(
                    f'{product_input.name}に対して有効な実金テンプレートが見つかりません。設定情報が正しく設定されていることをご確認ください。')

            template_path, output_path = self.path_info(
                product_input, form_no)
            output_path = before_ext.sub(
                '_' + '_'.join([
                    time_helper.strftime(
                        product_input.contract_date, r'%Y%m%d'
                    ),
                    product_input.customer_name.strip(),
                    product_input.product_kv['ファイル名用住所'],
                    product_input.name.strip()
                ]),
                output_path,
                1)

            return (template_path, output_path)

        logger.info('--- 実金シートの生成 ---')
        logger.info('')
        builders = []
        for idx, product_input in enumerate(product_inputs):
            builder = jikkin_sheet.output_jikkin_sheet(
                product_input,
                *jikkin_path(product_input),
                self.config
            )
            logger.info(
                f'  {idx + 1}列目 {product_input.name} -> {builder.src_jikkin_path}')
            builders.append(builder)
        logger.info('')

        logger.info('--- 計算処理 ---')
        logger.info('')
        logger.info(
            '  以下のファイルをExcelアプリケーションで開きます。出力が正しいことを確認し、必ず上書き保存してください。')
        logger.info('')
        for b in builders:
            logger.info(f'    {b.src_jikkin_path}')
        logger.info('')

        # excelを開く
        for b in builders:
            if platform.system() == 'Windows':
                # windowsで確認したところ、excelをセーブして閉じるまで
                # この呼出はブロックしてくれたので
                subprocess.run(f'"{b.src_jikkin_path}"', shell=True)
            # macでの処理
            elif platform.system() == 'Darwin':
                subprocess.Popen(["open", "-a", "LibreOffice",  f'{b.src_jikkin_path}',])

        logger.info('開いているすべてのExcelファイルを保存した後、Enterキーを押してください。')
        sys.stdin.flush()
        input()

        products = [builder.run() for builder in builders]

        # シートごとの出力

        # NOTE: 合計を出力する条件が難しいのでコメントでも残しておく。
        # 合計の出力は請求書が2枚以上発行される場合であるが、2054,9054については
        # 直前の列とセットでドキュメントを出力するため、単純にproductsの長さを測るだけ
        # では請求書の合計枚数と一致しない場合がある。

        def count_product(products: Iterable[Product]) -> int:
            """
            2054, 9054を考慮した実際の商品の数を計算する
            """
            return functools.reduce(
                lambda s, a: s
                if a[0] == "70N" and a[1] in {'2054', '9054'}
                else s + 1,
                more_itertools.pairwise(itertools.chain(
                    map(lambda p: p.name, products), (None, ))),
                0
            )

        if (count_product(products) >= 2):
            self._gen_gokei(products)
        if (len(products[0].joint_guarantors) > 0 and count_product(products) >= 2):
            self._gen_kashitsuke_gokei(products)

        # すべての商品の連帯保証人が等しくなるべき。
        if not all(p.joint_guarantors == products[0].joint_guarantors for p in products):
            raise RuntimeError('連帯保証人欄に不正があります。連帯保証人に関する項目はすべての列で同じ値を指定してください')

        self._gen_rentaihosho(products)
        self._gen_shinkokusho(products[0])

        """
        chacoを含む入力データかどうかをチェックする
        """
        def check_chacot(product_inputs:ProductInput):
            for idx, product_input in enumerate(product_inputs):
                if 'Chacot' in product_input.name:
                    chacot_flg[0] = 1
                    chacot_flg[1] += 1 
                    chacot_flg[2] = idx -1
                    chacot_flg[3] = idx
                    
        check_chacot(product_inputs)

        # 商品ごとの出力
        for product, next_product in more_itertools.pairwise(itertools.chain(products, (None, ))):
            for form_no, *_ in self.config.get_product_doc_info(product).items():

                assert type(form_no) == int

                src, dest = self.path_info(product.product_input, form_no)

                # NOTE: form_noごとに文書を作成する。特殊な処理が必要なものについては
                # self._gen_???? のような命名で専用の出力処理を作成してある。
                # 特殊な処理が必要ない場合はこのifの分岐の最後でデフォルトの置換ロジックを
                # 呼び出すような実装になっている。なので、特殊な出力をする(self._gen_????)
                # をした場合はcontinue文を実行して後続のデフォルトの置換ロジックまで計算が
                # 走らないようにする必要がある。

                if form_no in output_root_form_no:
                    # 直下に出力するものについては商品ごとの出力をする必要がないのでcontinue
                    continue

                elif form_no == 1:  # 請求書
                    # 70n and (2054 or 9054) は、2054 or 9054の請求書を作成し、70n分は出力しない
                    if next_product is not None and next_product.name in {'2054', '9054'}:
                        if product.name == '70N':
                            # 作成しているのはproductのではなく、next_productであることに注意
                            self._gen_bill(next_product, product)
                            continue
                        else:
                            RuntimeError('2054, 9054の直前の列は70Nである必要があります')
                    elif product.name in {'2054', '9054'}:
                        continue  # 2054, 9054は直前のループによって作成されているため、スキップ
                    else:
                        self._gen_bill(product)
                        continue

                elif form_no == 2:  # 依頼書
                    self._gen_iraisho(product)
                    continue

                elif form_no == 3:  # 申告書
                    # 申告書は契約者と連帯保証人ごとに出力するのでスキップ
                    continue

                elif form_no == 4:  # 金消
                    self._gen_kinsho(product)
                    continue

                elif form_no in {5, 6, 24, 25}:  # 事前説明書
                    self._gen_jizen(product)
                    continue

                elif form_no in {7, 8, 9}:  # 実金シート
                    continue

                elif form_no == 10:  # 大阪シート
                    self._gen_osaka(product)
                    continue

                elif form_no in {11, 12}:  # DEED
                    # 70n and (2054 or 9054) は2054 or 9054のdeedの書類を作成し、70n分は出力しない
                    if next_product is not None and next_product.name in {'2054', '9054'}:
                        if product.name == '70N':
                            # 作成しているのはproductのではなく、next_productであることに注意
                            self._gen_deed(next_product, product)
                            continue
                        else:
                            RuntimeError('2054, 9054の直前の列は70Nである必要があります')
                    elif product.name in {'2054', '9054'}:
                        continue  # 2054, 9054は直前のループによって作成されているため、スキップ
                    else:
                        self._gen_deed(product)
                        continue

                elif form_no in {13, 14}:  # note
                    # 後続のデフォルトの置換ロジックを用いるため、他の分岐とは違いcontinue
                    # しない
                    dest = before_ext.sub('_' + '_'.join([
                        time_helper.strftime(product.contract_date, r'%Y%m%d'),
                        product.customer_name,
                        product.product_input.product_kv['ファイル名用住所'],
                        product.name
                    ]), dest, 1
                    )
                elif form_no == 19:  # 金消Chacot
                    self._gen_kinsho_chacot(products, product)
                    continue

                elif form_no in {20, 21}:
                    src, dest = self.path_info(product.product_input, form_no)
                    if product.product_input.product_kv['担保物件所有者区分（Ｃｈａｃｏｔ）'] == '別':
                        src = before_ext.sub(
                            f'{"別"}',
                            src,
                            count=1) 
                    dest = before_ext.sub('_' + '_'.join([
                        time_helper.strftime(product.contract_date, r'%Y%m%d'),
                        product.customer_name,
                        product.product_input.product_kv['ファイル名用住所'],
                        product.name
                    ]), dest, 1
                    )
                elif form_no == 23:
                    if not (product.product_input.product_kv['担保物件所有者区分（Ｃｈａｃｏｔ）'] == '別' and product.state == 'GA'):
                        continue
                    dest = before_ext.sub('_' + '_'.join([
                        time_helper.strftime(product.contract_date, r'%Y%m%d'),
                        product.customer_name,
                        product.product_input.product_kv['ファイル名用住所'],
                        product.name
                    ]), dest, 1
                    )

                elif form_no == 26:  # 依頼書chacot
                    self._gen_iraisho_chacot(product)
                    continue
                else:
                    dest = before_ext.sub('_' + '_'.join([
                        time_helper.strftime(product.contract_date, r'%Y%m%d'),
                        product.customer_name,
                        product.product_input.product_kv['ファイル名用住所'],
                        product.name
                    ]), dest, 1
                    )

                os.makedirs(os.path.dirname(dest), exist_ok=True)
                replace.replace(
                    src, dest, ChainMap(
                        self.config.get_kv_for_product(
                            product.name, form_no, product.state),
                        product.table_kv, product.product_input
                    ))
                word_to_pdf_2_pages_per_sheet(dest, os.path.splitext(dest)[0] + ".pdf")

                logging_output(src, dest)

    def _gen_gokei(self, products: List[Product]):
        """
        様式番号17 の 請求書合計を出力する
        """
        if len(products) == 0:
            raise RuntimeError('商品が一つも存在しません。')

        form_no = 17
        src, dest = self.path_info(products[0].product_input, form_no)
        wb = openpyxl.load_workbook(src)

        ws = cast(Worksheet, wb['請求書合計'])

        customer_name = products[0].product_input.get('顧客名') \
            or products[0].product_input.get('法人名')
        if customer_name is None:
            raise RuntimeError('顧客名、または法人名が有効な値ではありません')
        ws['A1'].value = customer_name

        # 各商品の行のスタイル
        style_src_cells = ws['A3:B3'][0]

        # 2054,9054の直前にある70Nを取り除いた商品のリスト
        filtered_product = functools.reduce(
            lambda acc, product_tuple: acc
            if product_tuple[0].name == "70N" and product_tuple[1] is not None and product_tuple[1].name in {'2054', '9054'}
            else [*acc, product_tuple[0]],
            more_itertools.pairwise(itertools.chain(products, (None, ))),
            []
        )

        for product, row in zip(filtered_product, ws.iter_rows(min_row=3, max_row=1000)):
            # 5行目以降の処理は、現在の行に合計があるのでひとつ下にずらす
            if row[0].row >= 5:
                for current, below in zip(*ws.iter_rows(row[0].row, row[0].row + 1, 1, 2)):
                    below.value = current.value
                    xl_helper.copy_style(current, below)
            key_address_for_property_address = 'Ｐｒｏｐｅｒｔｙ　Ａｄｄ'
            key_total_billing_yen = '請求額合計（円）'
            row[0].value = product.product_input.get(
                key_address_for_property_address)
            row[1].value = product.product_input.get(key_total_billing_yen)

            if row[0].value is None:
                raise RuntimeError(
                    f'{key_address_for_property_address}に値がありません。')

            if row[1].value is None:
                raise RuntimeError(f'{key_total_billing_yen}に値がありません。')

            # 4行目以降スタイルのコピペ
            if row[0].row >= 4:
                for s, d in zip(style_src_cells, row):
                    xl_helper.copy_style(s, d)

        ws.cell(row[0].row + 1, 2).value = f"=SUM(B3:B{row[0].row})"

        dest = before_ext.sub('_' + '_'.join(
            [time_helper.strftime(products[0].contract_date, r'%Y%m%d'),
             products[0].customer_name,
             ]), dest, 1)
        os.makedirs(os.path.dirname(dest), exist_ok=True)
        wb.save(dest)
        replace.replace(dest, dest, ChainMap(
            products[0].table_kv,
            products[0].product_input.product_kv,
            self.config.get_kv_for_product(products[0].name, form_no, products[0].state)))
        logging_output(src, dest)

        return wb

    def _gen_iraisho(self, product: Product):
        """
        依頼書の作成
        """
        form_no = 2

        src, dest = self.path_info(
            product.product_input, form_no)
        dest = before_ext.sub('_' + '_'.join(
            [time_helper.strftime(product.contract_date, r'%Y%m%d'),
             product.customer_name,
             product.product_input.product_kv['ファイル名用住所'],
             ]), dest, 1
        )
        src = before_ext.sub('1'
                             if len(product.joint_guarantors) <= 1
                             else '4',
                             src)
        os.makedirs(os.path.dirname(dest), exist_ok=True)

        docx = Document(src)
        key = '連帯保証人住所'
        keyword = f'{replace.keyword_quoter}{key}{replace.keyword_quoter}'

        table = more_itertools.first_true(
            docx.tables,
            pred=lambda table: any(
                keyword in run.text
                for run in docx_helper.run_in_table(table))
        )

        for joint_guarantor, run in zip(
            itertools.chain(product.joint_guarantors,
                            itertools.repeat(JointGuarantor())),
            (run for run in docx_helper.run_in_table(table)
             if keyword in run.text)
        ):
            run.text = replace.replace_template_string(
                run.text,
                {key:  joint_guarantor.address},
                src)

        # 使わない行をテーブルから削除
        # 連帯保証人は最低1行残す
        row_height = 2
        for row in table.rows[row_height * (max(1, len(product.joint_guarantors)) + 1):]:
            table._tbl.remove(row._tr)

        docx.save(dest)
        replace.replace(dest, dest, ChainMap(
            product.table_kv,
            product.product_input,
            self.config.get_kv_for_product(
                product.name, form_no, product.state)
        ))
        word_to_pdf_2_pages_per_sheet(dest, os.path.splitext(dest)[0] + ".pdf")
        logging_output(src, dest)

    def _gen_iraisho_chacot(self, product: Product):
        """
        依頼書の作成
        """
        form_no = 26

        src, dest = self.path_info(
            product.product_input, form_no)
        dest = before_ext.sub('_' + '_'.join(
            [time_helper.strftime(product.contract_date, r'%Y%m%d'),
             product.customer_name,
             product.product_input.product_kv['ファイル名用住所'],
             ]), dest, 1
        )
        src = before_ext.sub('1_Chacot'
                             if len(product.joint_guarantors) <= 1
                             else '4_Chacot',
                             src)
        os.makedirs(os.path.dirname(dest), exist_ok=True)

        docx = Document(src)
        key = '連帯保証人住所'
        keyword = f'{replace.keyword_quoter}{key}{replace.keyword_quoter}'

        table = more_itertools.first_true(
            docx.tables,
            pred=lambda table: any(
                keyword in run.text
                for run in docx_helper.run_in_table(table))
        )

        for joint_guarantor, run in zip(
            itertools.chain(product.joint_guarantors,
                            itertools.repeat(JointGuarantor())),
            (run for run in docx_helper.run_in_table(table)
             if keyword in run.text)
        ):
            run.text = replace.replace_template_string(
                run.text,
                {key:  joint_guarantor.address},
                src)

        # 使わない行をテーブルから削除
        # 連帯保証人は最低1行残す
        row_height = 2
        for row in table.rows[row_height * (max(1, len(product.joint_guarantors)) + 1):]:
            table._tbl.remove(row._tr)

        docx.save(dest)
        replace.replace(dest, dest, ChainMap(
            product.table_kv,
            product.product_input,
            self.config.get_kv_for_product(
                product.name, form_no, product.state)
        ))
        word_to_pdf_2_pages_per_sheet(dest, os.path.splitext(dest)[0] + ".pdf")
        logging_output(src, dest)

    def _gen_kashitsuke_gokei(self, products: List[Product]):
        """
        様式番号18 の 貸付金額合計(社内用)を出力する
        """
        assert len(products) >= 2

        product = products[0]

        form_no = 18
        src, dest = self.path_info(product.product_input, form_no)
        dest = before_ext.sub('_' + '_'.join(
            [time_helper.strftime(product.contract_date, r'%Y%m%d'),
             product.customer_name,
             ]), dest, 1
        )

        wb = openpyxl.load_workbook(src)
        ws = cast(Worksheet, wb['Sheet1'])

        # 各商品の行のスタイル
        style_src_cells = ws['A2:E2'][0]

        for row_idx, item_idx, product, row in zip(
                itertools.count(2),
                itertools.count(1),
                products,
                ws.iter_rows(min_row=2, max_row=1000)
        ):

            def set_cell(cell: Cell, key: str):
                value = product.product_input.get(key)
                if value is None:
                    logger.debug(f'"{key}"は置換できません')
                cell.value = value

            # 行の値を書き込み
            row[0].value = item_idx
            set_cell(row[1], '商品区分')
            set_cell(row[2], '担保明細－物件名')
            set_cell(row[3], '貸付元本額（円）')

            # 3行目以降スタイルのコピペ
            if row_idx >= 3:
                for s, d in zip(style_src_cells, row):
                    xl_helper.copy_style(s, d)

        ws.cell(row_idx + 1, 1).value = None
        ws.cell(row_idx + 1, 2).value = None
        ws.cell(row_idx + 1, 3).value = None
        ws.cell(row_idx + 1, 4).value = f"=SUM(D2:D{row_idx})"
        ws.cell(row_idx + 1, 4).number_format = \
            ws.cell(row_idx, 4).number_format
        ws.cell(row_idx + 1, 4).font = \
            copy.copy(ws.cell(row_idx, 4).font)
        ws.cell(row_idx + 1, 4).alignment = \
            copy.copy(ws.cell(row_idx, 4).alignment)

        os.makedirs(os.path.dirname(dest), exist_ok=True)
        wb.save(dest)
        replace.replace(dest, dest, ChainMap(
            product.table_kv,
            product.product_input,
            self.config.get_kv_for_product(
                product.name, form_no, product.state)
        ))
        logging_output(src, dest)

        return wb

    def _gen_bill(self, product: Product, product_70n: Optional[Product] = None):
        """
        請求書の出力をする

        product_70n が Noneの場合通常の置換、Noneではない場合は2054, 9054用の特殊処理を実行する
        """

        form_no = 1

        src, dest = self.path_info(product.product_input, form_no)
        dest = before_ext.sub('_' + '_'.join(
            [time_helper.strftime(product.contract_date, r'%Y%m%d'),
             product.customer_name,
             product.product_input.product_kv['ファイル名用住所'],
             ]), dest, 1
        )
        os.makedirs(os.path.dirname(dest), exist_ok=True)

        new_src = src
        if product.name in {'2054', '9054'}:
            # 債権2の請求書は特殊な作りになっている

            key_product_name_master = "商品名マスタ"
            key_bill_yen = "請求額（円）"
            all_product_kv = ChainMap(
                product.table_kv,
                product.product_input,
                self.config.get_kv_for_product(
                    product.name, form_no, product.state)
            )

            all_product_70n_kv = ChainMap(
                product_70n.table_kv,
                product_70n.product_input,
                self.config.get_kv_for_product(
                    product_70n.name, form_no, product.state)
            )

            docx = Document(src)
            product_name_masters = [run for run in docx_helper.all_runs(
                docx) if key_product_name_master in run.text]
            bills = [run for run in docx_helper.all_runs(
                docx) if key_bill_yen in run.text]

            values = {
                key_product_name_master + "１": all_product_70n_kv[key_product_name_master],
                key_product_name_master + "２": all_product_kv[key_product_name_master],
                key_bill_yen + "１": product_70n.product_input[key_bill_yen],
                key_bill_yen + "２": product.product_input[key_bill_yen]
            }

            for run in itertools.chain(product_name_masters, bills):

                if key_product_name_master in run.text or key_bill_yen in run.text:
                    run.text = replace.replace_template_string(
                        run.text, values, src)

            docx.save(dest)
            new_src = dest



        # 辞書内情報確認
        chain_map = ChainMap(
            product.table_kv,
            product.product_input,
            self.config.get_kv_for_product(product.name, form_no, product.state)
        )

        # for key, value in self.config.get_kv_for_product(product.name, form_no, product.state).items():
        #     if value is not None:
        #         print(f"{key}: {value}")

        for key, value in vars(product.product_input).items():
            if value is not None:
                print(f"{key}: {value}")






        replace.replace(new_src, dest, ChainMap(
            product.table_kv,
            product.product_input,
            self.config.get_kv_for_product(
                product.name, form_no, product.state)
        ))
        word_to_pdf_2_pages_per_sheet(dest, os.path.splitext(dest)[0] + ".pdf")
        logging_output(src, dest)

    def _gen_osaka(self, product: Product):
        """
        大阪シートを作成する
        """

        form_no = 10

        src, dest = self.path_info(product.product_input, form_no)
        dest = before_ext.sub(
            '_' + '_'.join([
                time_helper.strftime(product.contract_date, r'%Y%m%d'),
                product.customer_name,
                product.product_input.product_kv['ファイル名用住所'],
                product.name,
            ]),
            dest, 1
        )

        def gen_row(row: Tuple[Any, ...]
                    # 実行及び返済日, 年月日, 約弁回, 手数料等, 後払利息, 返済元本, 約定元本, 利用可能額, 計算元本積数, 日数
                    ):
            # 日付, 回数, 利率, 弁済額, 元金, 利息, 残高
            return lambda idx: (
                row[1].value,
                row[0].value,
                product.jikkin_kv.get('約定利率'),
                f"=E{idx}+F{idx}",
                row[5].value,
                row[4].value,
                f"=G{idx-1} - E{idx}"
            )

        src_wb = openpyxl.load_workbook(product.jikkin_path, data_only=True)

        jikkin_sheet = cast(Worksheet, src_wb['実金'])

        osaka_sheet_wb = openpyxl.load_workbook(src)
        osaka_sheet = cast(Worksheet, osaka_sheet_wb['Sheet1'])

        zipped_jikkin_osaka = list(zip(
            # 実行及び返済日,年月日,手数料等,後払利息,返済元本, 約定元本, 利用可能額, 計算元本積数,日数
            jikkin_sheet.iter_rows(23, jikkin_sheet.max_row, 1, 10),
            # 適当に多めに1000行取る。zipするので1000行読むわけでもない。
            osaka_sheet.iter_rows(10, 1000, 1, 7)
        ))
        max_rows = 0
        # NOTE: 以下はtable_kvから取得してはいけない。Excelでは参照先が空でも参照元にデフォルト値の0が入ってしまうから。
        osaka_sheet['G5'].value = product.jikkin_kv['最終弁済時ＬＴＶ']

        osaka_sheet['G9'].value = product.table_kv['貸付元本額（￥）']

        for jikkin_row, osaka_row in zipped_jikkin_osaka:
            if '合計' in str(jikkin_row[0].value):
                osaka_row[1].value = '合計'

                osaka_row[3].value = f'=SUM(D{zipped_jikkin_osaka[0][1][0].row}:D{osaka_row[0].row - 1})'
                osaka_row[3].number_format = zipped_jikkin_osaka[1][0][3].number_format

                osaka_row[4].value = f'=SUM(E{zipped_jikkin_osaka[0][1][0].row}:E{osaka_row[0].row - 1})'
                osaka_row[4].number_format = zipped_jikkin_osaka[1][0][4].number_format

                osaka_row[5].value = f'=SUM(F{zipped_jikkin_osaka[0][1][0].row}:F{osaka_row[0].row - 1})'
                osaka_row[5].number_format = zipped_jikkin_osaka[1][0][5].number_format
                #最大行数の取得
                max_rows = osaka_row[0].row + 1

                for cell in osaka_row[1:6]:
                    cell.fill = PatternFill(
                        patternType='solid', fgColor='FFFFFF00')
                    side = Side(style='thin', color='000000')
                    cell.border = Border(
                        top=side, left=side, right=side, bottom=side)

            else:
                values = gen_row(jikkin_row)(osaka_row[0].row)
                for value, cell, style_src_cell in zip(values, osaka_row, zipped_jikkin_osaka[0][1]):
                    cell.value = value
                    xl_helper.copy_style(style_src_cell, cell)

        #スタイルが1000行分適用されるため不要行の削除
        for osaka_row in reversed(range(max_rows, 1000)):
            #行を削除
            osaka_sheet.delete_rows(osaka_row)

        os.makedirs(os.path.dirname(dest), exist_ok=True)
        osaka_sheet_wb.save(dest)
        replace.replace(dest, dest, ChainMap(
            product.table_kv,
            product.product_input,
            self.config.get_kv_for_product(
                product.name, form_no, product.state)
        ))
        logging_output(src, dest)

        # ExcelファイルをPDFに変換
        pdf_dest = os.path.splitext(dest)[0] + '.pdf'
        excel_to_pdf(dest, pdf_dest)

    def _gen_shinkokusho(self, product: Product):
        """
        申告書を作成する
        """

        form_no = 3

        for joint_guarantor in product.joint_guarantors:
            src, dest = self.path_info(
                product.product_input, form_no)
            dest = before_ext.sub('_' + '_'.join(
                [time_helper.strftime(product.contract_date, r'%Y%m%d'),
                 product.customer_name,
                 joint_guarantor.name
                 ]), dest, 1
            )

            os.makedirs(os.path.dirname(dest), exist_ok=True)
            replace.replace(src, dest, ChainMap(
                {'郵便番号': joint_guarantor.postal_code,
                    '顧客住所': joint_guarantor.address},
                product.table_kv,
                product.product_input,
                self.config.get_kv_for_product(
                    product.name, form_no, product.state)
            ))
            word_to_pdf_2_pages_per_sheet(dest, os.path.splitext(dest)[0] + ".pdf")
            logging_output(src, dest)

        # 個人契約者向けに出力
        if product.is_personal:
            src, dest = self.path_info(
                product.product_input, form_no)
            dest = before_ext.sub('_' + '_'.join(
                [time_helper.strftime(product.contract_date, r'%Y%m%d'),
                 product.customer_name,
                 ]), dest, 1
            )
            os.makedirs(os.path.dirname(dest), exist_ok=True)
            replace.replace(src, dest, ChainMap(
                product.table_kv,
                product.product_input,
                self.config.get_kv_for_product(
                    product.name, form_no, product.state)
            ))
            word_to_pdf_2_pages_per_sheet(dest, os.path.splitext(dest)[0] + ".pdf")
            logging_output(src, dest)

    def _gen_jizen(self, product: Product):
        """
        貸付契約事前説明書(アモチ有り無し)を作成する
        """
        form_no, _ = more_itertools.first_true(
            self.config.get_product_doc_info(product).items(),
            default=(None, None),
            pred=lambda x: x[0] in {5, 6, 24, 25},
        )

        src, dest = self.path_info(
            product.product_input, form_no)
        dest = before_ext.sub('_' + '_'.join(
            [time_helper.strftime(product.contract_date, r'%Y%m%d'),
             product.customer_name,
             product.product_input.product_kv['ファイル名用住所'],
             product.name
             ]), dest, 1
        )
        os.makedirs(os.path.dirname(dest), exist_ok=True)
        # Chacotかどうかでテンプレートを選ぶ
        if form_no in {24, 25}:
            src = before_ext.sub(
                'Chacot',
                src,
                1)
        if form_no in {24, 25} and product.product_input.product_kv['担保物件所有者区分（Ｃｈａｃｏｔ）'] == '別':
            src = before_ext.sub(
                f'{"別"}',
                src,
                count=1) 
        docx = Document(src)
        key = '連帯保証人住所'
        keyword = f'{replace.keyword_quoter}{key}{replace.keyword_quoter}'

        table = more_itertools.first_true(
            docx.tables,
            pred=lambda table: any(
                keyword in run.text
                for run in docx_helper.run_in_table(table))
        )

        for joint_guarantor, run in zip(
            itertools.chain(
                product.joint_guarantors,
                itertools.repeat(JointGuarantor())
            ),
            (run for run in docx_helper.run_in_table(table)
             if keyword in run.text)
        ):
            run.text = replace.replace_template_string(
                run.text,
                {key: joint_guarantor.address},
                src)

        # 使わない行をテーブルから削除
        # 連帯保証人は最低1行残す
        row_height = 3
        for row in table.rows[row_height * (max(1, len(product.joint_guarantors)) + 1):]:
            table._tbl.remove(row._tr)

        docx.save(dest)
        new_src = dest

        replace.replace(new_src, dest, ChainMap(
            product.table_kv,
            product.product_input,
            self.config.get_kv_for_product(
                product.name, form_no, product.state)
        ))
        word_to_pdf_2_pages_per_sheet(dest, os.path.splitext(dest)[0] + ".pdf")
        logging_output(src, dest)

    def _gen_deed(self, product: Product, product_70n: Optional[Product] = None):
        """
        DEEDの出力をする

        product_70n が Noneの場合通常の置換、Noneではない場合は2054, 9054用の特殊処理を実行する
        """

        form_no, _ = more_itertools.first_true(
            self.config.get_product_doc_info(product).items(),
            default=(None, None),
            pred=lambda x: x[0] in {11, 12, 21},
        )

        self.spell_number(product.product_input, product.state)

        src, dest = self.path_info(
            product.product_input, form_no)
        dest = before_ext.sub('_' + '_'.join(
            [time_helper.strftime(product.contract_date, r'%Y%m%d'),
             product.customer_name,
             product.product_input.product_kv['ファイル名用住所'],
             ]), dest, 1
        )
        os.makedirs(os.path.dirname(dest), exist_ok=True)

        new_src = src
        if form_no == 12:
            docx = Document(src)

            values = {
                "金消契約日１": product_70n.contract_date,
                "金消契約日２": product.contract_date,
                "最終弁済日１": product_70n.product_input['最終弁済日'],
                "最終弁済日２": product.product_input['最終弁済日']
            }

            for run in docx_helper.all_runs(docx):

                if "金消契約日" in run.text or "最終弁済日" in run.text:
                    run.text = replace.replace_template_string(
                        run.text, values, src)

            docx.save(dest)
            new_src = dest

        replace.replace(new_src, dest, ChainMap(
            product.table_kv,
            product.product_input,
            self.config.get_kv_for_product(
                product.name, form_no, product.state)
        ))
        word_to_pdf_2_pages_per_sheet(dest, os.path.splitext(dest)[0] + ".pdf")
        logging_output(src, dest)

    def _gen_rentaihosho(self, products: Iterable[Product]):
        """
        様式 15, 16を連帯保証人の数だけ出力する

        NOTE: 入力シートから複数列入力をする場合でも、連帯保証人はどの列でも同じ値が指定されている
        ことを期待している。
        """

        product = next(iter(products))

        for form_no in {15, 16}:

            for joint_guarantor in product.joint_guarantors:
                src, dst = self.path_info(product.product_input, form_no)

                output_file_path=before_ext.sub('_' + '_'.join(
                    [time_helper.strftime(product.contract_date, r'%Y%m%d'),
                        product.customer_name,
                        joint_guarantor.name
                        ]
                    ), dst, 1)
                replace.replace(
                    input_file_path=src,
                    output_file_path=output_file_path,
                    replace_dict=ChainMap(
                        {  # product_input は未加工のsplit前のものが入っているので、こちらでうわがき
                            "連帯保証人名": joint_guarantor.name,
                            "連帯保証人住所": joint_guarantor.address,
                        },
                        product.product_input,
                        product.table_kv,
                        self.config.get_kv_for_product(
                            product.name, form_no, product.state)
                    ))
                word_to_pdf_2_pages_per_sheet(output_file_path, os.path.splitext(output_file_path)[0] + ".pdf")

    def _gen_kinsho(self, product: Product):
        form_no = 4
        kinsho_kv = self.config.get_kv_for_product(
            product.name, form_no, product.state)
        src, dest = self.path_info(product.product_input, form_no)

        key = "連帯保証人住所"

        template_key = f'{replace.keyword_quoter}{key}{replace.keyword_quoter}'
        # 連帯保証人の数に合わせてテンプレートを選ぶ
        src = before_ext.sub(
            '1' if len(product.joint_guarantors) <= 1 else '4',
            src,
            1)
        docx = Document(src)
        table = more_itertools.first_true(docx.tables,  pred=lambda table: any(
            template_key in run.text for run in docx_helper.run_in_table(table)))

        table_runs = [run
                      for run in docx_helper.run_in_table(table)
                      if template_key in run.text
                      ]

        # 連帯保証人住所を出現順に置換する。
        for joint_guarantor, run in zip(
            itertools.chain(
                product.joint_guarantors,
                itertools.repeat(JointGuarantor())),
            table_runs,
        ):

            if run is None:
                raise RuntimeError(
                    f'{src}に含まれている{key}のテーブルの数が足りません。テンプレート側のテーブルの行数を増やしてから再度実行してください。')
            if joint_guarantor is not None:
                run.text = replace.replace_template_string(
                    run.text,
                    {key: joint_guarantor.address},
                    src)
            else:
                run.text = ''  # タグを消す。

        # 余ったtableの行を削除
        # 連帯保証人は最低1行残す
        row_height = 2
        for row in table.rows[row_height * (max(1, len(product.joint_guarantors)) + 1):]:
            table._tbl.remove(row._tr)

        # 一旦書き出す

        dest = before_ext.sub(

            '_' + '_'.join([
                time_helper.strftime(product.contract_date, r'%Y%m%d'),
                product.customer_name,
                product.product_input.product_kv['ファイル名用住所'],
                product.name,
            ]),
            dest, 1


        )
        os.makedirs(os.path.dirname(dest), exist_ok=True)
        docx.save(dest)
        replace.replace(dest, dest, ChainMap(
            kinsho_kv,
            product.table_kv,
            product.product_input
        ))
        word_to_pdf_2_pages_per_sheet(dest, os.path.splitext(dest)[0] + ".pdf")
        logging_output(src, dest)

    def _gen_kinsho_chacot(self, products, product: Product):
        form_no = 19
        kinsho_kv = self.config.get_kv_for_product(
            product.name, form_no, product.state)
        src, dest = self.path_info(product.product_input, form_no)

        key = "連帯保証人住所"

        template_key = f'{replace.keyword_quoter}{key}{replace.keyword_quoter}'
        # 連帯保証人の数に合わせてテンプレートを選ぶ
        src = before_ext.sub(
            'Chacot',
            src,
            1)
        # Chacot別判定
        if product.product_input.product_kv['担保物件所有者区分（Ｃｈａｃｏｔ）'] == '別':
            src = before_ext.sub(
                f'{"別"}',
                src,
                count=1) 
        docx = Document(src)
        table = more_itertools.first_true(docx.tables,  pred=lambda table: any(
            template_key in run.text for run in docx_helper.run_in_table(table)))

        table_runs = [run
                      for run in docx_helper.run_in_table(table)
                      if template_key in run.text
                      ]

        # 連帯保証人住所を出現順に置換する。
        for joint_guarantor, run in zip(
            itertools.chain(
                product.joint_guarantors,
                itertools.repeat(JointGuarantor())),
            table_runs,
        ):

            if run is None:
                raise RuntimeError(
                    f'{src}に含まれている{key}のテーブルの数が足りません。テンプレート側のテーブルの行数を増やしてから再度実行してください。')
            if joint_guarantor is not None:
                run.text = replace.replace_template_string(
                    run.text,
                    {key: joint_guarantor.address},
                    src)
            else:
                run.text = ''  # タグを消す。
        
        """
        表のコピーを行う
        """
        def copy_table_after(table, paragraph):
            tbl, p = table._tbl, paragraph._p
            new_tbl = copy.deepcopy(tbl)
            p.addnext(new_tbl)

        target_p = 0

        for i, table in enumerate(docx.tables):
            for cell in table._cells:
                if '＜別紙＞ 物件目録-●物件番号●（本件不動産）' in cell.text:
                    target_p = i
                    break
        if target_p == 0:
            RuntimeError('テンプレートファイル内の文字列が変更されたため、正しく物件目録が出力されていない可能性があります。')

        tables_cnt = len(docx.tables)
        template = docx.tables[target_p]
        paragraph = docx.paragraphs[149]

        if chacot_flg[0] == 1 and chacot_flg[1] == 1 and chacot_flg[2] > 0: 
            x = 0
            while x < chacot_flg[2]:
                copy_table_after(template, paragraph)
                x += 1

        # 余ったtableの行を削除
        # 連帯保証人は最低1行残す
        row_height = 2
        for row in table.rows[row_height * (max(1, len(product.joint_guarantors)) + 1):]:
            table._tbl.remove(row._tr)

        # 一旦書き出す

        dest = before_ext.sub(

            '_' + '_'.join([
                time_helper.strftime(product.contract_date, r'%Y%m%d'),
                product.customer_name,
                product.product_input.product_kv['ファイル名用住所'],
                product.name,
            ]),
            dest, 1

        )
        cnt = 0
        # 置換前にテーブルの置換を行う
        for i, table in enumerate(docx.tables):
            for cell in table._cells:
                if '＜別紙＞ 物件目録-●物件番号●（本件不動産）' in cell.text:
                    products[cnt].product_input.product_kv['物件番号'] = str(cnt + 1)
                    for run in docx_helper.run_in_table(table):
                        if isinstance(run.text, str):
                            replaced_str = replace.replace_template_string(
                                run.text, ChainMap(
                                    kinsho_kv,
                                    products[cnt].table_kv,
                                    products[cnt].product_input
                                ), dest)
                            if run.text is not replaced_str:
                                run.text = replaced_str
                    cnt += 1

        if cnt == 0:
            logger.debug('テンプレートファイル内の文字列が変更されたため、正しく物件目録が出力されていない可能性があります。')

        product.product_input.product_kv['物件番号'] = str(cnt + 1)
        os.makedirs(os.path.dirname(dest), exist_ok=True)
        docx.save(dest)
        replace.replace(dest, dest, ChainMap(
            kinsho_kv,
            product.table_kv,
            product.product_input
        ))
        word_to_pdf_2_pages_per_sheet(dest, os.path.splitext(dest)[0] + ".pdf")
        logging_output(src, dest)



    def path_info(self, product_input: ProductInput, form_no: int) -> Tuple[str, str]:
        """
        商品と様式からシステム上のファイルパスを計算して返却する

        戻り値は、(テンプレートのパス, 未加工の出力先のパス名)
        """

        if not isinstance(product_input, ProductInput):
            raise TypeError()

        template_file = self.config.get_template_filename(
            product_input, form_no)

        contractor_dir = '_'.join([
            strftime(product_input.contract_date, r"%Y%m%d"),
            product_input.customer_name.strip()]
        )

        property_dir = '' \
            if form_no in output_root_form_no \
            else '_'.join([
                strftime(product_input.contract_date, r"%Y%m%d"),
                product_input.customer_name.strip(),
                product_input.product_kv['ファイル名用住所']
            ])

        return (
            os.path.normpath(os.path.join(
                self.template_root_path, template_file)),
            os.path.normpath(os.path.join(
                self.output_root_path,
                contractor_dir,
                property_dir,
                template_file
            )))


    def spell_number(self, product_input: ProductInput, state: str):
            if product_input.name == 'コーポレート50':
                if state == 'HI' and product_input.product_kv['貸付元本額（英）'] != '':
                    num = product_input.product_kv['貸付元本額（英）']
                    product_input.product_kv['貸付元本額（英）'] = self.gine_number(num)

    def gine_number(self, num):
        dic_num = {
            '0' : '',
            '1' : 'One',
            '2' : 'Two',
            '3' : 'Three',
            '4' : 'Four',
            '5' : 'Five',
            '6' : 'Six',
            '7' : 'Seven',
            '8' : 'Eight',
            '9' : 'Nine',
            '10' : 'Ten',
            '11' : 'Eleven',
            '12' : 'Twelve',
            '13' : 'Thirteen',
            '14' : 'Fourteen',
            '15' : 'Fifteen',
            '16' : 'Sixteen',
            '17' : 'Seventeen',
            '18' : 'Eighteen',
            '19' : 'Nineteen',
            '20' : 'Twenty ',
            '30' : 'Thirty ',
            '40' : 'Forty ',
            '50' : 'Fifty ',
            '60' : 'Sixty ',
            '70' : 'Seventy ',
            '80' : 'Eighty ',
            '90' : 'Ninety ',
            '00' : ''
        }

        def get_digit(dic_num: dict, num: str):
            dec_place = num.find(".")
            if dec_place > 0:
                num = num[:dec_place -1]

            count = 1
            digit = { 1:'', 2:' Thousand ', 3:' Million ', 4:' Billion ', 5:' Trillon ', 6:'', 7:''  }
            dollars = '';
            while num != '':
                temp = get_hundreds(dic_num, num[-3:])
                if temp != '':
                    dollars = temp + digit[count] + dollars
                if len(num) > 3:
                    num = num[:len(num)-3] 
                else:
                    num = ''
                count += 1 
            return dollars

        def get_hundreds(dic_num:dict, num: str):
            num = num.zfill(3)
            if len(num) % 100 == 0:
                return dic_num[num[0]] + ' Hundred '
            if int(num[0]) == 0:
                return get_tens(dic_num, num[1:]) 
            else:
                return dic_num[num[0]] + ' Hundred ' + get_tens(dic_num, num[1:]) 

        def get_tens(dic_num: dict, num: str):
            if int(num) <= 20:
                return dic_num[num]
            elif int(num) % 10 == 0:
                return dic_num[num]
            else:
                return dic_num[num[0] + '0'] +  dic_num[num[1]]
        
        num = str(num).strip()
        number = get_digit(dic_num, num)
        return number
