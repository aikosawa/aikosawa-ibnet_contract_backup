from docx.document import Document
from docx.section import Sections
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from itertools import chain
from typing import Iterable, List


def run_in_paragraphs(paragraphs: List[Paragraph]) -> Iterable[Run]:
    return (run for p in paragraphs for run in p.runs)


def run_in_table(table: Table) -> Iterable[Run]:
    return (run
            for row in table.rows
            for cell in row.cells
            for run in chain(run_in_paragraphs(cell.paragraphs), run_in_tables(cell.tables)))


def run_in_tables(tables: List[Table]) -> Iterable[Run]:
    return (run
            for table in tables
            for run in run_in_table(table))


def run_in_paragraphs_and_tables(obj) -> Iterable[Run]:
    return chain(run_in_paragraphs(obj.paragraphs), run_in_tables(obj.tables))


def run_in_sections(sections: Sections):
    return chain(*(chain(run_in_paragraphs_and_tables(section.header),
                         run_in_paragraphs_and_tables(section.footer))
                   for section in sections))


def all_runs(doc: Document):
    return chain(run_in_paragraphs_and_tables(doc),
                 run_in_sections(doc.sections))
